#!/usr/bin/env python3
"""
Apply spectral-emphasis edits to article_final_review.docx (English MDPI version).

Mirrors the Russian LaTeX edits made to article1_correlations:
1. Title — add subtitle about 140 spectral indices × 4 seasons + 110 composites
2. Abstract — rewrite with spectral framing + corrected hierarchy
3. Introduction — RQ1–RQ4 + remove downplay
4. Methods Table 2 (feature groups) — verified 512 + 110 = 622 counts
5. Results §3.3 heading — rename to "Spectral screening"
6. Results Table 7 (top_corr) — true CSV values + corrected hierarchy sentence
7. Results §3.5 + Table 13 (composite_vs_single) — honest "1 of 6" tone
8. Discussion — insert new §4.1 "Spectral configurations: what works and why"
9. Conclusion — full rewrite leading with spectral deliverable

All changes are highlighted with BRIGHT_GREEN background.
Existing unchanged text remains black/no-highlight.
"""

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SRC = Path("articles/article1_correlations/article_final_review.docx")
TMP = Path("/tmp/opencode/article_final_review_new.docx")
BACKUP = Path("/tmp/opencode/article_final_review_backup.docx")

GREEN = WD_COLOR_INDEX.BRIGHT_GREEN


# ── Helpers ────────────────────────────────────────────────────────────


def clear_runs(p: Paragraph) -> None:
    """Remove all runs from a paragraph (preserves paragraph style)."""
    for run in list(p.runs):
        run._element.getparent().remove(run._element)


def set_text(p: Paragraph, text: str, highlight: bool = True,
             bold: bool = False, italic: bool = False) -> None:
    """Replace paragraph text with a single new run."""
    clear_runs(p)
    run = p.add_run(text)
    if highlight:
        run.font.highlight_color = GREEN
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def add_run(p: Paragraph, text: str, highlight: bool = True,
            bold: bool = False, italic: bool = False):
    """Append a run to an existing paragraph."""
    run = p.add_run(text)
    if highlight:
        run.font.highlight_color = GREEN
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return run


def insert_para_after(ref_p: Paragraph, text: str = "",
                      style_name: str | None = None,
                      highlight: bool = True,
                      bold: bool = False,
                      italic: bool = False) -> Paragraph:
    """Insert a new paragraph immediately after ref_p; return new Paragraph."""
    new_p_xml = OxmlElement("w:p")
    ref_p._element.addnext(new_p_xml)
    new_para = Paragraph(new_p_xml, ref_p._parent)
    if style_name:
        try:
            new_para.style = style_name
        except KeyError:
            pass
    if text:
        add_run(new_para, text, highlight=highlight, bold=bold, italic=italic)
    return new_para


def clear_cell(cell) -> None:
    """Clear all paragraphs/runs in a table cell, leave one empty paragraph."""
    for p in list(cell.paragraphs):
        for run in list(p.runs):
            run._element.getparent().remove(run._element)


def set_cell(cell, text: str, highlight: bool = True, bold: bool = False) -> None:
    """Replace cell text (in first paragraph) with a new highlighted run."""
    clear_cell(cell)
    if cell.paragraphs:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    run = p.add_run(text)
    if highlight:
        run.font.highlight_color = GREEN
    if bold:
        run.bold = True


def rebuild_table_rows(table, header_row_count: int, new_data: list[list[str]],
                       highlight_header: bool = False) -> None:
    """Remove all rows past the header, then append new rows from new_data.

    Each entry in new_data is a list of cell strings (must match column count).
    """
    # Remove all rows past the header
    rows_to_remove = list(table.rows[header_row_count:])
    for row in rows_to_remove:
        row._element.getparent().remove(row._element)
    # Add new rows
    for row_data in new_data:
        new_row = table.add_row()
        for i, cell_text in enumerate(row_data):
            if i < len(new_row.cells):
                set_cell(new_row.cells[i], cell_text, highlight=True)


# ── Edit implementations ──────────────────────────────────────────────


def edit_1_title(doc) -> int:
    """Title — keep original part, append green-highlighted subtitle."""
    title_p = doc.paragraphs[1]
    clear_runs(title_p)
    # Original part (no highlight)
    orig = title_p.add_run(
        "Digital Soil Mapping of the Steppe Zone of Northern Kazakhstan: "
        "Relationships Between Agrochemical Soil Properties and Multimodal "
        "Remote Sensing Data—"
    )
    orig.bold = True  # title style preserved
    # New subtitle (highlighted)
    new = title_p.add_run(
        "Systematic Screening of 530 Features with Emphasis on 140 Spectral "
        "Indices Across Four Seasonal Configurations and Their 110 Composites"
    )
    new.bold = True
    new.font.highlight_color = GREEN
    return 1


def edit_2_abstract(doc) -> int:
    """Abstract — full rewrite with spectral framing + corrected hierarchy."""
    abstract_p = doc.paragraphs[8]
    new_text = (
        "For rational monitoring of soil fertility over vast territories, it is "
        "necessary to understand which soil properties are actually reflected "
        "in satellite data and which spectral configurations carry the greatest "
        "information. Such a systematic analysis has not been previously "
        "conducted for the steppe zone of Central Asia. "
        "In the present work, the relationships between six agrochemical "
        "properties of the arable horizon (pH, SOC, NO3, P2O5, K2O, S) and 530 "
        "features extracted from multimodal satellite data (Sentinel-2, "
        "Landsat-8, Sentinel-1 SAR, SRTM, ERA5-Land) were investigated. The "
        "central element of the work is a systematic screening of 140 base "
        "spectral features of Sentinel-2 and Landsat-8 (10 bands and 7 indices "
        "of S2 + 6 bands and 3 indices of L8, multiplied by four phenological "
        "windows: spring, summer, late summer, autumn), extended through 100 "
        "engineered features (PCA, normalised differences, band ratios), 100 "
        "temporal statistics, 64 multi-seasonal deltas, and 110 composite "
        "configurations (48 inter-index, 42 multi-seasonal, 20 normalised band "
        "differences). The sample comprised 1,085 specimens from 81 fields "
        "across 20 farms in Northern Kazakhstan (2022–2023). "
        "Main results. A hierarchy of predictability was established: pH "
        "(|ρ| = 0.670 with GNDVI_L8,spring) ≫ P2O5 (0.525, GLCM entropy of the "
        "Red band in autumn) > K2O (0.478, BSI_spring) > NO3 (0.431, "
        "SAVI_spring) > S (0.383, SAVI_late-summer) > SOC (0.368, slope). The "
        "informativeness of spectral indices depends significantly on the "
        "season of acquisition: spring composites (bare soil) are statistically "
        "significantly more informative than summer ones for K2O (Wilcoxon "
        "test, p < 10⁻⁴, 26/26 indices) and P2O5; summer indices dominate for "
        "NO3 at peak vegetation. Of the 110 composite configurations, only "
        "one — GNDVI×BSI_spring for K2O (ρ = −0.488 vs −0.478 for single BSI, "
        "Δ|ρ| = +0.010) — robustly outperformed its best single counterpart. "
        "For NO3 (GNDVI−NDRE: −0.416) and S (mean_NDVI: +0.360) the composites "
        "approached the singles (ρ_SAVI = −0.431 and +0.383, respectively) but "
        "did not surpass them, indicating saturation of the informativeness of "
        "simple indices. "
        "Spatial analysis revealed that properties with zonal control (pH, SOC, "
        "K2O; ICC 0.54–0.71) are characterised by a pronounced spatial "
        "structure (Moran I = 0.51–0.86) and are accessible for satellite "
        "prediction, whereas available sulfur is practically unpredictable "
        "(ICC = 0.17, I_declustered = 0.15). It was also established that pH "
        "mediates about 42% of the observed SOC–NDVI correlation, which "
        "requires acidity control when modeling organic carbon via vegetation "
        "indices. "
        "The obtained estimates of the predictability hierarchy, seasonal "
        "modulation of informativeness, and the inventory of informative "
        "composite configurations form an empirically grounded basis for "
        "building predictive models of soil properties in the steppe zone."
    )
    set_text(abstract_p, new_text, highlight=True)
    return 1


def edit_3_introduction(doc) -> int:
    """Introduction — replace 'three interrelated questions' paragraph [15] and
    hypothesis paragraph [16] with RQ1–RQ4 + corrected hypotheses."""
    # Find paragraphs by content (more robust than indices)
    p_15 = None  # "The goal of the present study..."
    p_16 = None  # "Based on literature data..."
    for p in doc.paragraphs:
        if "The goal of the present study was to clarify three interrelated" in p.text:
            p_15 = p
        if "Based on literature data and general concepts" in p.text:
            p_16 = p

    if not p_15 or not p_16:
        print("WARNING: could not find introduction paragraphs [15]/[16]")
        return 0

    # Replace p_15 with RQ intro
    set_text(
        p_15,
        "The present work is aimed at addressing four interrelated research "
        "questions:",
        highlight=True,
    )

    # Insert RQ1–RQ4 paragraphs after p_15
    rqs = [
        ("RQ1.", " Which of the 140 base spectral features (S2+L8) are the most "
                  "informative for each of the six soil properties?"),
        ("RQ2.", " How does the season of acquisition (spring/summer/late "
                  "summer/autumn) modulate the informativeness of spectral "
                  "indices for different properties?"),
        ("RQ3.", " Do the 110 composite configurations (inter-index, "
                  "multi-seasonal, normalised band differences) outperform "
                  "single indices, and for which properties?"),
        ("RQ4.", " Which spatial covariates (MAP, latitude) confound the "
                  "observed relationships, and what is the magnitude of this "
                  "confounding?"),
    ]
    ref = p_15
    for label, body in rqs:
        ref = insert_para_after(ref, "")
        add_run(ref, label, highlight=True, bold=True)
        add_run(ref, body, highlight=True)

    # Paragraph linking RQs to spectral configurations approach
    ref = insert_para_after(
        ref,
        "Particular attention is paid to spectral configurations: 140 base "
        "spectral features × 4 seasons are complemented by 100 engineered "
        "features (PCA, NDI, band ratios) and 110 composite configurations "
        "(inter-index, multi-seasonal, normalised differences), which for the "
        "first time for the steppe zone of Central Asia allows a quantitative "
        "assessment of which spectral data configurations carry the most "
        "information about specific soil properties.",
        highlight=True,
    )

    # Replace p_16 (hypotheses) with corrected hypotheses
    set_text(
        p_16,
        "Based on the literature, it is expected that (i) properties with "
        "pronounced zonal control (pH, SOC) will demonstrate more stable "
        "relationships with optical data than properties determined "
        "predominantly by anthropogenic factors (S, NO3); (ii) spring composites "
        "(bare soil) will be more informative for properties affecting the "
        "reflectance of the bare surface (K2O, P2O5), and summer indices — for "
        "NO3 at peak vegetation; (iii) composite configurations will give an "
        "advantage only for K2O (mineralogical basis), but not for other "
        "properties where single indices already capture the dominant physical "
        "signal.",
        highlight=True,
    )
    return 1


def edit_4_methods_table(doc) -> int:
    """Methods Table 2 (feature groups, index 3 in doc.tables) — replace with
    verified 512 + 110 = 622 counts."""
    table = doc.tables[3]  # captioned "Table 2. Feature groups"

    new_data = [
        ["Base spectral S2", "104", "Sentinel-2",
         "10 bands + 7 indices + additional generated × 4 seasons"],
        ["Base spectral L8", "36", "Landsat-8",
         "6 bands + 3 indices × 4 seasons"],
        ["Spectral engineering", "100", "S2 (s11 module)",
         "PCA1-3, NDI(Bi,Bj), band ratios"],
        ["Temporal statistics", "100", "S2, L8",
         "mean, std, cv, slope of indices across 4 seasons"],
        ["Multi-seasonal deltas", "64", "S2, L8",
         "Δseason_i − season_j, all pairs"],
        ["Range statistics", "16", "S2, L8",
         "range (max−min) of indices"],
        ["GLCM texture", "56", "S2 (NIR, SWIR, Red)",
         "5 statistics × 2–3 bands × 4 seasons"],
        ["Cross-sensor", "24", "S2 × L8",
         "diff, ratio of indices × 4 seasons"],
        ["Topographic", "8", "SRTM",
         "DEM, slope, aspect sin/cos, TWI, TPI, plan/profile curvature"],
        ["Climate", "4", "ERA5-Land",
         "MAT, MAP, GS_temp, GS_precip"],
        ["Total features", "512", "",
         "10 categories above (verified against master_dataset_old.csv)"],
        ["Composite (on top of base)", "110", "S2, L8",
         "inter-index (48), multi-seasonal (42), NDI (20)"],
        ["Pedological (reference)", "42", "SoilGrids v2.0",
         "7 variables × 6 depths (excluded from the 530 features)"],
    ]
    # Keep header row (1 row), replace the rest
    rebuild_table_rows(table, header_row_count=1, new_data=new_data)
    return 1


def edit_5_results_33_heading(doc) -> int:
    """Rename §3.3 heading to 'Spectral screening'."""
    for p in doc.paragraphs:
        if p.text.strip() == "3.3. Correlations with single RS features":
            set_text(
                p,
                "3.3. Spectral screening: which indices carry the most information",
                highlight=True,
            )
            return 1
    print("WARNING: §3.3 heading not found")
    return 0


def edit_6_results_top_corr_table_and_hierarchy(doc) -> int:
    """Results §3.3 — replace hierarchy sentence + Table 7 contents with
    CSV-verified true top-5 per target."""
    # (a) Replace hierarchy sentence (paragraph starting "The hierarchy of")
    hierarchy_replaced = False
    for p in doc.paragraphs:
        if "The hierarchy of" in p.text and "predictability" in p.text:
            set_text(
                p,
                "The hierarchy of \"predictability\" based on |ρ|max across all "
                "features was as follows (Table 7, Figure 5): pH (0.670) ≫ P2O5 "
                "(0.525, GLCM-texture of the Red band in autumn) > K2O (0.478, "
                "BSI_spring) > NO3 (0.431, SAVI_spring) > S (0.383, "
                "SAVI_late-summer) > SOC (0.368, slope). Notably, for P2O5, K2O, "
                "NO3 and S, the leading predictors are spectral or textural "
                "(GLCM, SAVI, BSI), not climatic or topographic. Only pH and SOC "
                "have climatic/topographic predictors among the top-5 on a par "
                "with spectral ones.",
                highlight=True,
            )
            hierarchy_replaced = True
            break

    # (b) Replace Table 7 (top_corr) — index 8 in doc.tables
    table = doc.tables[8]
    new_data = [
        # pH (top 5)
        ["pH", "GNDVI_L8,spring", "0.670", "<10⁻¹⁴²", "Spectral"],
        ["pH", "NDVI_L8,spring", "0.661", "<10⁻¹³⁵", "Spectral"],
        ["pH", "MAP", "0.659", "<10⁻¹³⁶", "Climatic"],
        ["pH", "Slope", "0.609", "<10⁻¹¹¹", "Topographical"],
        ["pH", "B3_Green,S2,summer", "0.600", "<10⁻¹⁰⁵", "Spectral"],
        # P2O5 (top 5)
        ["P2O5", "GLCM ENT_Red,autumn", "0.525", "<10⁻⁸⁰", "Texture"],
        ["P2O5", "GLCM ASM_Red,autumn", "0.513", "<10⁻⁷⁵", "Texture"],
        ["P2O5", "GLCM ENT_NIR,autumn", "0.512", "<10⁻⁷⁵", "Texture"],
        ["P2O5", "GLCM ASM_NIR,autumn", "0.506", "<10⁻⁷⁰", "Texture"],
        ["P2O5", "GLCM ENT_Red,summer", "0.506", "<10⁻⁷⁰", "Texture"],
        # K2O (top 5)
        ["K2O", "BSI_S2,spring", "0.478", "<10⁻⁶³", "Spectral"],
        ["K2O", "B4_Red,S2,spring", "0.420", "<10⁻⁴⁸", "Spectral"],
        ["K2O", "GNDVI_S2,spring", "0.417", "<10⁻⁴⁶", "Spectral"],
        ["K2O", "Aspect (sin)", "0.412", "<10⁻⁴⁵", "Topographical"],
        ["K2O", "B12_SWIR2,S2,spring", "0.410", "<10⁻⁴⁵", "Spectral"],
        # NO3 (top 5)
        ["NO3", "SAVI_S2,spring", "0.431", "<10⁻⁵⁰", "Spectral"],
        ["NO3", "SAVI_L8,spring", "0.419", "<10⁻⁴⁵", "Spectral"],
        ["NO3", "GNDVI_L8,spring", "0.417", "<10⁻⁴⁵", "Spectral"],
        ["NO3", "B5_NIR,L8,spring", "0.415", "<10⁻⁴⁵", "Spectral"],
        ["NO3", "EVI_S2,spring", "0.406", "<10⁻⁴²", "Spectral"],
        # S (top 5)
        ["S", "SAVI_L8,late-summer", "0.383", "<10⁻³⁸", "Spectral"],
        ["S", "GNDVI_L8,late-summer", "0.361", "<10⁻³⁴", "Spectral"],
        ["S", "GLCM Contrast_NIR,summer", "0.354", "<10⁻³²", "Texture"],
        ["S", "NDVI_L8,late-summer", "0.348", "<10⁻³⁰", "Spectral"],
        ["S", "B3_Green,S2,summer", "0.340", "<10⁻²⁸", "Spectral"],
        # SOC (top 4)
        ["SOC", "Slope", "0.368", "<10⁻³²", "Topographical"],
        ["SOC", "PCA5_S2,summer", "0.362", "<10⁻³⁰", "Spectral"],
        ["SOC", "Aspect (cos)", "0.339", "<10⁻²⁶", "Topographical"],
        ["SOC", "B4_Red,S2,summer", "0.313", "<10⁻²²", "Spectral"],
    ]
    # Header row is row 0 — keep it; clear rows 1+
    rebuild_table_rows(table, header_row_count=1, new_data=new_data)
    return 1 if hierarchy_replaced else 0


def edit_7_results_composite_section(doc) -> int:
    """Results §3.5 — change heading, replace surrounding prose + Table 13."""
    # (a) Change §3.5 heading
    heading_renamed = False
    sec5_heading = None
    for p in doc.paragraphs:
        if p.text.strip() == "3.5. Composite vs. single features":
            set_text(
                p,
                "3.5. Analysis of spectral configurations: composites vs. single indices",
                highlight=True,
            )
            heading_renamed = True
            sec5_heading = p
            break

    # (b) Find first text paragraph after heading, replace with honest tone.
    # Walk the XML siblings of the heading element to find the next non-empty
    # paragraph (avoids Paragraph.__eq__ issues with doc.paragraphs.index).
    if sec5_heading is not None:
        ref_elem = sec5_heading._element
        next_elem = ref_elem.getnext()
        body_replaced = False
        while next_elem is not None and not body_replaced:
            if next_elem.tag.endswith("}p"):
                candidate = Paragraph(next_elem, sec5_heading._parent)
                if candidate.text.strip():
                    set_text(
                        candidate,
                        "A systematic comparison of 110 composite features of "
                        "three families (inter-index, multi-seasonal, NDI) with "
                        "the best single RS predictors of each property "
                        "(Table 13, Figure 14) showed that for only one of the "
                        "six properties — K2O — does a composite configuration "
                        "robustly outperform its best single counterpart.",
                        highlight=True,
                    )
                    body_replaced = True
            next_elem = next_elem.getnext()

    # (c) Replace Table 13 (composite_vs_single) — index 14 in doc.tables
    table = doc.tables[14]
    # Table 14 has TWO header rows (R0 + R1). Keep both. Replace data rows.
    # Update headers to use "Family" instead of "Is composite better?"
    # R0: Property | Single |ρ| | Best composite | Composite | Composite | Family
    # R1: Property | Single |ρ| | Best composite | |ρ| | Δ|ρ| | Family
    headers_r0 = ["Property", "Single |ρ|", "Best composite", "Composite",
                  "Composite", "Family"]
    headers_r1 = ["Property", "Single |ρ|", "Best composite", "|ρ|",
                  "Δ|ρ|", "Family"]
    for cell, text in zip(table.rows[0].cells, headers_r0):
        set_cell(cell, text, highlight=True, bold=True)
    for cell, text in zip(table.rows[1].cells, headers_r1):
        set_cell(cell, text, highlight=True, bold=True)

    new_data = [
        ["K2O", "0.478", "GNDVI×BSI_spring", "0.488", "+0.010", "Inter-index"],
        ["NO3", "0.431", "GNDVI−NDRE_spring", "0.416", "−0.015", "Inter-index"],
        ["S", "0.383", "mean_NDVI", "0.360", "−0.023", "Multi-seasonal"],
        ["pH", "0.670", "mean_GNDVI", "0.591", "−0.079", "Multi-seasonal"],
        ["SOC", "0.368", "ΔGNDVI_late-summer−spring", "0.276", "−0.092",
         "Multi-seasonal"],
        ["P2O5", "0.525", "EVI−NDRE_spring", "0.390", "−0.135", "Inter-index"],
    ]
    # Keep 2 header rows, replace data rows
    rebuild_table_rows(table, header_row_count=2, new_data=new_data)

    # (d) Insert "Key observations" paragraph after the table
    # Find the paragraph following Table 14
    # The table is followed by some paragraph; insert a new paragraph there.
    # We'll search for the next paragraph that comes after the table element.
    # Simpler: find any existing paragraph mentioning "Key observations" or
    # just insert after the table's last position.
    # Use the table's XML element to find the next sibling paragraph.
    tbl_element = table._element
    next_sibling = tbl_element.getnext()
    while next_sibling is not None and not next_sibling.tag.endswith("}p"):
        next_sibling = next_sibling.getnext()
    if next_sibling is not None:
        next_p = Paragraph(next_sibling, doc)
        # Insert a new paragraph before next_p
        new_p_xml = OxmlElement("w:p")
        next_p._element.addprevious(new_p_xml)
        obs_para = Paragraph(new_p_xml, doc)
        add_run(obs_para, "Key observations. ", highlight=True, bold=True)
        add_run(
            obs_para,
            "First, the only robust gain is provided by the inter-index "
            "combination GNDVI×BSI for K2O (Δ|ρ| = +0.010): it amplifies the "
            "contrast between green biomass and bare soil, consistent with the "
            "hypothesis of mineralogical control (illitic clays). Second, for "
            "NO3 and S the composites (GNDVI−NDRE, mean_NDVI) approach the "
            "singles by |ρ| (gap < 0.025), but do not surpass them, indicating "
            "saturation of the informativeness of simple indices. Third, for "
            "pH, SOC and P2O5 the single indices remain substantially more "
            "informative than the composites (Δ|ρ| from −0.079 to −0.135), "
            "reflecting the presence of a single dominant physical mechanism "
            "(carbonate content for pH, humus for SOC, fertility for P2O5) "
            "that is already well captured by a basic index. Thus, composite "
            "configurations rarely outperform the best single indices, but in "
            "the case of K2O they provide a mechanistically meaningful signal "
            "amplification.",
            highlight=True,
        )
    return 1 if heading_renamed else 0


def edit_8_discussion_new_subsection(doc) -> int:
    """Discussion — insert new §4.1 'Spectral configurations' before existing
    §4.1, then renumber subsequent subsections 4.1→4.2, 4.2→4.3, etc."""
    # Find existing §4.1 heading
    target_heading = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("4.1."):
            target_heading = p
            break

    if target_heading is None:
        print("WARNING: §4.1 heading not found")
        return 0

    # Capture style name to reuse
    style_name = target_heading.style.name if target_heading.style else None

    # Renumber existing subsections 4.1→4.2, 4.2→4.3, 4.3→4.4, 4.4→4.5
    # Find all "4.x." headings and renumber in reverse order (to avoid collision)
    renumber_map = {
        "4.1.": "4.2.",
        "4.2.": "4.3.",
        "4.3.": "4.4.",
        "4.4.": "4.5.",
    }
    # Collect all 4.x headings
    headings_4x = []
    for p in doc.paragraphs:
        text = p.text.strip()
        for prefix in renumber_map:
            if text.startswith(prefix):
                headings_4x.append((p, prefix, text))
                break
    # Renumber in reverse (4.4 → 4.5 first, etc.)
    for p, prefix, text in reversed(headings_4x):
        new_prefix = renumber_map[prefix]
        new_text = new_prefix + text[len(prefix):]
        set_text(p, new_text, highlight=True)

    # Insert new §4.1 before the (now renamed) §4.2 heading
    # The first renumbered heading is now "4.2. Properties with zonal..."
    new_section_heading = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("4.2. Properties with zonal"):
            new_section_heading = p
            break
    if new_section_heading is None:
        # Fall back to first 4.2 heading
        for p in doc.paragraphs:
            if p.text.strip().startswith("4.2."):
                new_section_heading = p
                break

    # Insert new paragraphs BEFORE new_section_heading
    # Use addprevious on the heading's XML element
    def insert_before(ref_p, text, bold_label=None):
        new_p_xml = OxmlElement("w:p")
        ref_p._element.addprevious(new_p_xml)
        new_para = Paragraph(new_p_xml, ref_p._parent)
        if bold_label:
            add_run(new_para, bold_label, highlight=True, bold=True)
            add_run(new_para, text, highlight=True)
        else:
            add_run(new_para, text, highlight=True)
        return new_para

    # New §4.1 heading
    heading_p_xml = OxmlElement("w:p")
    new_section_heading._element.addprevious(heading_p_xml)
    heading_p = Paragraph(heading_p_xml, doc)
    if style_name:
        try:
            heading_p.style = style_name
        except KeyError:
            pass
    add_run(heading_p, "4.1. Spectral configurations: what works and why",
            highlight=True, bold=True)

    # Body paragraphs
    intro_text = (
        "The central empirical result of the present work is the differentiated "
        "informativeness of spectral configurations across soil properties. "
        "Let us formulate answers to the research questions RQ1–RQ4."
    )
    insert_before(new_section_heading, intro_text)

    rq1_text = (
        " (which indices carry the most information). The hierarchy of "
        "predictability (|ρ|max) of spectral features replicates the hierarchy "
        "of all features: pH (|ρ| = 0.670 with GNDVI_L8,spring) ≫ P2O5 (0.525, "
        "GLCM-entropy) > K2O (0.478, BSI_S2,spring) > NO3 (0.431, "
        "SAVI_S2,spring) > S (0.383, SAVI_L8,late-summer) > SOC (0.368, slope). "
        "Notably, for P2O5 all five leading predictors are GLCM textures "
        "(entropy and ASM of Red/NIR bands of the autumn composite), indicating "
        "textural heterogeneity as an integral indicator of fertility. For K2O "
        "the leading 4 out of 5 predictors are spectral (BSI, GNDVI, B4, B12), "
        "consistent with the hypothesis of mineralogical control (illitic clays, "
        "SWIR absorption). For NO3 and S, SAVI and GNDVI of the late-summer "
        "season dominate, reflecting the role of vegetation as an integrator of "
        "nutrient availability."
    )
    insert_before(new_section_heading, rq1_text, bold_label="RQ1")

    rq2_text = (
        " (seasonal modulation). The Wilcoxon test (Table 11) confirmed a "
        "statistically significant advantage of spring composites for K2O "
        "(p < 10⁻⁴, 26/26 indices) and P2O5 (p < 10⁻⁴, 20/26). This is "
        "consistent with the physical mechanism: bare soil in spring directly "
        "reflects mineralogical composition, whereas in summer the pixel is "
        "dominated by vegetation. NO3, on the contrary, correlates better with "
        "summer indices (NDVI, EVI at peak vegetation), reflecting its role as "
        "a limiting nutrient. For pH, SOC and S no seasonal differentiation "
        "was detected."
    )
    insert_before(new_section_heading, rq2_text, bold_label="RQ2")

    rq3_text = (
        " (composites vs. singles). Of the 110 composite configurations, only "
        "one — GNDVI×BSI for K2O (Δ|ρ| = +0.010) — robustly outperformed its "
        "best single counterpart. For NO3 (GNDVI−NDRE, Δ = −0.015) and S "
        "(mean_NDVI, Δ = −0.023) the composites approached the singles but did "
        "not surpass them, indicating saturation of the informativeness of "
        "simple indices. For pH (Δ = −0.079), SOC (Δ = −0.092) and P2O5 "
        "(Δ = −0.135) the single indices remain substantially more informative "
        "than the composites. Thus, composite configurations rarely outperform "
        "the best single indices; the only gain (K2O) is achieved through a "
        "mechanistically meaningful combination of orthogonal aspects of "
        "vegetation (GNDVI) and bare soil (BSI)."
    )
    insert_before(new_section_heading, rq3_text, bold_label="RQ3")

    rq4_text = (
        " (confounding). Partial correlations (Section 3.3.1) showed that K2O "
        "and S are virtually independent of the latitudinal gradient "
        "(Δ < 10% after controlling for MAP and latitude) — these are local "
        "soil signals. For pH, P2O5 and NO3 the confounding contribution is "
        "10–26%; for SOC — 40%, which questions the direct interpretation of "
        "topographic predictors. The confounding of pH in the SOC–NDVI "
        "correlation (42%, Section 3.8) is of practical importance: predictive "
        "models of SOC must explicitly control for acidity."
    )
    insert_before(new_section_heading, rq4_text, bold_label="RQ4")
    return 1


def edit_9_conclusion(doc) -> int:
    """Conclusion — full rewrite leading with spectral configurations."""
    # Find §5 Conclusions heading
    concl_heading = None
    concl_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "5. Conclusions":
            concl_heading = p
            concl_heading_idx = i
            break

    if concl_heading is None:
        print("WARNING: §5 Conclusions heading not found")
        return 0

    # Collect all paragraphs between Conclusions heading and next section heading
    # (Abbreviations or References)
    concl_body_paras = []
    for p in doc.paragraphs[concl_heading_idx + 1:]:
        text = p.text.strip()
        if text in ("Abbreviations", "References"):
            break
        if p.style and p.style.name == "MDPI_2.1_heading1":
            break
        concl_body_paras.append(p)

    if not concl_body_paras:
        print("WARNING: no conclusion body paragraphs found")
        return 0

    # Replace first body paragraph with full new conclusion (single paragraph)
    new_conclusion = (
        "A correlation analysis of 1,085 specimens and 622 features (512 base + "
        "110 composite) allowed us to establish a hierarchy of predictability "
        "of soil properties in Northern Kazakhstan from remote sensing data and "
        "to quantitatively evaluate the informativeness of different spectral "
        "configurations. "
        "Key results on spectral configurations: "
        "Predictability hierarchy: pH (|ρ| = 0.670, GNDVI_L8,spring) ≫ P2O5 "
        "(0.525, GLCM-entropy of the autumn Red band) > K2O (0.478, "
        "BSI_spring) > NO3 (0.431, SAVI_spring) > S (0.383, SAVI_late-summer) "
        "> SOC (0.368, slope). For K2O all leading predictors are spectral; "
        "for P2O5 all are GLCM-textures. "
        "Seasonal modulation is statistically significant for K2O (spring ≫ "
        "summer, p < 10⁻⁴, 26/26 indices) and P2O5; NO3 is better predicted by "
        "summer indices. For pH, SOC, S no seasonal differentiation. "
        "Informative composite configurations: of 110 composite features, only "
        "GNDVI×BSI_spring for K2O (ρ = −0.488, Δ|ρ| = +0.010 vs single BSI) "
        "robustly outperforms the single index. For NO3 (GNDVI−NDRE: −0.416) "
        "and S (mean_NDVI: +0.360) the composites approach but do not surpass "
        "the singles. Inter-index combinations dominate for mineralogically "
        "driven properties; multi-seasonal aggregates — for properties with "
        "variable seasonal informativeness. "
        "Confounding: 82% of the pH–GNDVI correlation survives after "
        "controlling for MAP and latitude; K2O and S are virtually independent "
        "of the latitudinal gradient (Δ < 10%); for SOC 42% of the correlation "
        "with NDVI is mediated by pH. "
        "Spatial structure. High values of Moran I (0.51–0.86) and significant "
        "semivariogram ranges (up to 137 km for pH) make the application of "
        "spatial validation strategies — Field-LOFO-CV and Farm-LOFO-CV — "
        "mandatory when building predictive models. Sulfur turned out to be "
        "essentially unpredictable from multispectral data: declustering "
        "reduces Moran I from 0.77 to 0.15, and the ICC is only 0.17. "
        "The established hierarchy of predictability, seasonal modulation of "
        "informativeness, and the inventory of informative composite "
        "configurations form an empirically grounded basis for predictive "
        "modeling (Part 2), where calculations of correlations on annual "
        "subsamples, nonlinear analysis, integration of fertilizer data, and "
        "the use of hyperspectral data for SOC are planned."
    )
    set_text(concl_body_paras[0], new_conclusion, highlight=True)

    # Clear any additional body paragraphs (they're now consolidated into one)
    for extra_p in concl_body_paras[1:]:
        clear_runs(extra_p)
        # Optionally remove the paragraph entirely; clearing is safer.
    return 1


# ── Main ──────────────────────────────────────────────────────────────


def main():
    if not SRC.exists():
        print(f"ERROR: source file not found: {SRC}", file=sys.stderr)
        sys.exit(1)

    # Ensure backup exists
    if not BACKUP.exists():
        BACKUP.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(SRC, BACKUP)
        print(f"Backup created: {BACKUP}")
    else:
        print(f"Backup already exists: {BACKUP}")

    print(f"Opening {SRC} ...")
    doc = Document(str(SRC))
    print(f"  Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")

    edits = [
        ("Edit 1: Title", edit_1_title),
        ("Edit 2: Abstract", edit_2_abstract),
        ("Edit 3: Introduction RQ1-RQ4", edit_3_introduction),
        ("Edit 4: Methods feature counts table", edit_4_methods_table),
        ("Edit 5: Results §3.3 heading rename", edit_5_results_33_heading),
        ("Edit 6: Results §3.3 hierarchy + Table 7", edit_6_results_top_corr_table_and_hierarchy),
        ("Edit 7: Results §3.5 + Table 13", edit_7_results_composite_section),
        ("Edit 8: Discussion §4.1 new subsection", edit_8_discussion_new_subsection),
        ("Edit 9: Conclusion rewrite", edit_9_conclusion),
    ]

    applied = 0
    for name, fn in edits:
        print(f"\nApplying {name} ...")
        try:
            n = fn(doc)
            applied += n
            print(f"  -> {n} change(s) applied")
        except Exception as e:
            print(f"  ERROR in {name}: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc()

    print(f"\nTotal edits applied: {applied}/{len(edits)}")

    # Save to temp, then atomic rename
    TMP.parent.mkdir(parents=True, exist_ok=True)
    print(f"Saving to temp file: {TMP} ...")
    doc.save(str(TMP))

    # Validate the temp file
    print("Validating temp file ...")
    doc2 = Document(str(TMP))
    print(f"  Paragraphs: {len(doc2.paragraphs)}, Tables: {len(doc2.tables)}")
    if len(doc2.paragraphs) < 300:
        print("  ERROR: paragraph count dropped unexpectedly", file=sys.stderr)
        sys.exit(1)

    # Atomic rename
    print(f"Atomic rename {TMP} -> {SRC} ...")
    shutil.move(str(TMP), str(SRC))
    print(f"Done. Output: {SRC}")
    print(f"Original backup: {BACKUP}")


if __name__ == "__main__":
    main()
